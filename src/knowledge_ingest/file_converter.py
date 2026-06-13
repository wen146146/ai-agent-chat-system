import os
os.environ["PYTHONDONTWRITEBYTECODE"] = "1"

import re
import io
import csv
from pathlib import Path
from typing import Tuple
from dotenv import load_dotenv

load_dotenv(Path(__file__).parent.parent.parent / ".env")


def _detect_format(filename: str) -> str:
    ext = Path(filename).suffix.lower()
    format_map = {
        ".txt": "txt",
        ".md": "md",
        ".pdf": "pdf",
        ".docx": "docx",
        ".csv": "csv",
    }
    return format_map.get(ext, "txt")


def _convert_txt(content: str, filename: str) -> str:
    title = Path(filename).stem
    return f"# {title}\n\n{content}"


def _convert_md(content: str, filename: str) -> str:
    return content


def _convert_pdf(content_bytes: bytes, filename: str) -> str:
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(content_bytes))
        title = Path(filename).stem
        lines = [f"# {title}"]
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text and text.strip():
                lines.append(f"\n## 第{i+1}页\n\n{text.strip()}")
        return "\n".join(lines)
    except ImportError:
        return f"[错误] 缺少 pypdf 库，请运行: pip install pypdf"
    except Exception as e:
        return f"[错误] PDF转换失败: {str(e)}"


def _convert_docx(content_bytes: bytes, filename: str) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(content_bytes))
        title = Path(filename).stem
        lines = [f"# {title}"]
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                lines.append("")
                continue
            if para.style and para.style.name and para.style.name.startswith("Heading"):
                level = para.style.name.replace("Heading", "").strip()
                try:
                    level_num = int(level)
                except ValueError:
                    level_num = 2
                prefix = "#" * min(level_num, 4)
                lines.append(f"\n{prefix} {text}")
            else:
                lines.append(text)
        for table in doc.tables:
            lines.append("")
            lines.append("| " + " | ".join(["列"] * len(table.columns)) + " |")
            lines.append("| " + " | ".join(["---"] * len(table.columns)) + " |")
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                lines.append("| " + " | ".join(cells) + " |")
        return "\n".join(lines)
    except ImportError:
        return f"[错误] 缺少 python-docx 库，请运行: pip install python-docx"
    except Exception as e:
        return f"[错误] DOCX转换失败: {str(e)}"


def _convert_csv(content_str: str, filename: str) -> str:
    try:
        title = Path(filename).stem
        reader = csv.reader(content_str.splitlines())
        rows = list(reader)
        if not rows:
            return f"# {title}\n\n（空文件）"
        header = rows[0]
        lines = [
            f"# {title}",
            "",
            "| " + " | ".join(header) + " |",
            "| " + " | ".join(["---"] * len(header)) + " |",
        ]
        for row in rows[1:]:
            row_padded = row + [""] * (len(header) - len(row))
            lines.append("| " + " | ".join(row_padded) + " |")
        return "\n".join(lines)
    except Exception as e:
        return f"[错误] CSV转换失败: {str(e)}"


def convert_to_markdown(content, filename: str) -> str:
    """
    将各种格式文件统一转换为 Markdown 文本
    content: str(文本类) 或 bytes(二进制类)
    filename: 原始文件名
    """
    fmt = _detect_format(filename)

    if fmt == "txt":
        return _convert_txt(content, filename)
    elif fmt == "md":
        return _convert_md(content, filename)
    elif fmt == "pdf":
        content_bytes = content if isinstance(content, bytes) else content.encode("utf-8")
        return _convert_pdf(content_bytes, filename)
    elif fmt == "docx":
        content_bytes = content if isinstance(content, bytes) else content.encode("utf-8")
        return _convert_docx(content_bytes, filename)
    elif fmt == "csv":
        content_str = content if isinstance(content, str) else content.decode("utf-8")
        return _convert_csv(content_str, filename)
    else:
        return _convert_txt(str(content), filename)
